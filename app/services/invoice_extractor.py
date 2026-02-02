import re
from typing import Any

from app.services.llm_client import LLMClient, EXTRACT_PROMPT, REPAIR_PROMPT


class InvoiceExtractor:
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    async def extract(self, text: str) -> dict[str, Any]:
        normalized = self._normalize_text(text)
        payload = await self.llm_client.extract_json(EXTRACT_PROMPT, normalized)
        if payload is None:
            payload = {
                "supplier_name": None,
                "invoice_number": None,
                "invoice_date": None,
                "due_date": None,
                "incoterm": None,
                "currency": None,
                "total_value": None,
                "freight_cost": None,
                "insurance_cost": None,
                "line_items": [],
                "field_confidence": {},
                "confidence_score": 0.1,
                "warnings": ["LLM extraction unavailable"],
            }
        return payload


def detect_insurance_amount(text: str) -> float | None:
    patterns = [
        r"insurance\\s*(amount|premium|total)\\s*[:\\-]?\\s*([\\d,]+\\.?\\d*)",
        r"total\\s*insurance\\s*[:\\-]?\\s*([\\d,]+\\.?\\d*)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(match.lastindex)
            try:
                return float(value.replace(",", ""))
            except Exception:
                continue
    return None


async def extract_text_from_pdf(path: str) -> str:
    text = ""
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception:
        text = ""
    return text


async def extract_text_from_docx(path: str) -> str:
    try:
        import docx
    except Exception:
        return ""

    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


async def ocr_fallback(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""

    try:
        image = Image.open(path)
        return pytesseract.image_to_string(image)
    except Exception:
        return ""
